from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MD_PATH = DOCS / "PROJECT-TECHNICAL-DOCUMENT-ZH.md"
DOCX_PATH = DOCS / "PROJECT-TECHNICAL-DOCUMENT-ZH.docx"
DIAGRAM_PATH = DOCS / "project-technical-workflow.png"


SECTIONS = [
    (
        "1. 文件目的",
        [
            "本文件將 Subtitle Review Loop 整理為專案技術文件，說明需求來源、系統邊界、流程設計、功能模組、資料流、輸出格式、品質控管與部署方式。",
            "文件讀者包含專案維護者、課程影片製作人員、字幕校稿人員，以及想把此流程導入其他電腦或其他課程專案的使用者。",
        ],
    ),
    (
        "2. 專案背景與需求分析",
        [
            "教學影片與課程錄影常以 ASR 或 Whisper 產生草稿字幕，但初稿通常不能直接發布。實務上最容易出錯的地方不是 SRT 格式，而是專有名詞、斷句、口語清理、字幕位置與燒錄視覺效果。",
            "本專案採用 human-in-the-loop 設計：先以自動化工具提高效率，再透過瀏覽器校稿與樣片確認建立品質門檻。",
        ],
    ),
    (
        "3. 系統範圍",
        [
            "本專案負責字幕校稿流程、SRT 清理、燒錄設定、FFmpeg 燒字幕腳本、Windows 啟動檔與可攜 ZIP 工具包。",
            "本專案不內建雲端 ASR，也不把影片或字幕上傳到雲端；可攜工具包只負責校稿與燒錄設定，不直接轉錄或燒錄影片。",
        ],
    ),
    (
        "4. 整體工作流程",
        [
            "流程分為輸入收集、字幕產生或匯入、規則清理、字幕校稿、燒錄預覽、校稿包輸出、樣片確認與正式交付。",
            "每個階段都會留下可檢查的中間成果，避免發生錯誤時必須重做整支影片。",
        ],
    ),
    (
        "5. 功能模組設計",
        [
            "前端工作台由 src/subtitle-editor.html 實作，提供兩階段 UI：第一階段校稿字幕文字，第二階段設定與預覽燒錄字幕樣式。",
            "本機服務由 src/subtitle-editor-server.mjs 實作，支援靜態檔案、MP4 Range Request，以及 POST /api/save-review-package 儲存校稿包。",
            "燒字幕腳本由 scripts/burn-subtitles.mjs 實作，讀取 burn-settings.json 並轉成 FFmpeg force_style。",
            "可攜工具包由 scripts/build-portable-toolkit.mjs 建置，產生 input、output、啟動檔與使用說明。",
        ],
    ),
    (
        "6. 資料流與輸出格式",
        [
            "完整專案模式使用 workspace 作為本機素材與輸出位置；可攜工具包模式使用 input 與 output，降低非技術使用者操作門檻。",
            "主要輸出包含 media.edited.srt、burn-settings.json、burn-settings.ffmpeg-style.txt 與 export-manifest.json。",
        ],
    ),
    (
        "7. 品質保證與錯誤修正",
        [
            "品質檢查包含 SRT 解析、空 cue、時間軸重疊、專有名詞、rule file 套用、字幕位置、畫面遮擋、影片解析度與音訊保留。",
            "Runbook 已記錄 Windows UTF-8、PowerShell 亂碼、Whisper cp950、PowerShell heredoc、字幕過大與大型影片處理等問題。",
        ],
    ),
    (
        "8. 部署與操作",
        [
            "完整專案可透過 npm run open 啟動；Windows 使用者可雙擊 start-subtitle-editor.cmd。",
            "可攜工具包可透過 npm run build:toolkit 產生。解壓後，使用者將影片放入 input/media.mp4、字幕放入 input/media.srt，再雙擊啟動字幕校對工具.cmd。",
        ],
    ),
    (
        "9. 安全與隱私",
        [
            "workspace 與 dist 預設不提交 GitHub；影片、音訊、字幕與燒錄產物也被 .gitignore 排除。",
            "本機服務只監聽 127.0.0.1。分享工具包時應保持 input/output 為空，不附帶實際課程素材。",
        ],
    ),
    (
        "10. 後續擴充建議",
        [
            "建議新增閱讀速度與字數警示、空 cue/重疊 cue 視覺標示、一鍵產生 FFmpeg 樣片、中英雙語字幕欄位、glossary/typo map 匯入，以及工具包自訂輸出資料夾。",
        ],
    ),
]


TABLES = {
    "需求收斂": (
        ["類別", "技術需求"],
        [
            ["字幕校稿", "影片播放、cue 跳轉、搜尋、直接修改 SRT。"],
            ["兩階段流程", "第一階段校稿文字，第二階段設定燒錄樣式。"],
            ["規則管理", "清理規則由使用者提供，避免跨課程誤套。"],
            ["輸出標準", "產生 edited SRT、burn settings、FFmpeg style 與 manifest。"],
            ["Windows 友善", "雙擊啟動、UTF-8 防呆、Node.js 檢查。"],
            ["可攜分享", "建置 ZIP 工具包，其他電腦可解壓後使用。"],
        ],
    ),
    "模組清單": (
        ["檔案", "模組", "職責"],
        [
            ["src/subtitle-editor.html", "兩階段字幕校稿網頁", "影片同步、SRT 解析、cue 編輯、字幕樣式預覽與校稿包儲存。"],
            ["src/subtitle-editor-server.mjs", "本機 HTTP 服務", "靜態檔案、MP4 Range Request、POST /api/save-review-package。"],
            ["src/apply_subtitle_rules.mjs", "字幕清理腳本", "依使用者提供的 rule file 進行清理與報告輸出。"],
            ["scripts/burn-subtitles.mjs", "FFmpeg 燒字幕腳本", "讀取 burn-settings.json，轉換 force_style，支援樣片與正式輸出。"],
            ["scripts/build-portable-toolkit.mjs", "可攜工具包建置", "產生 input/output、雙擊啟動檔與使用說明。"],
        ],
    ),
    "輸出格式": (
        ["檔案", "說明", "用途"],
        [
            ["media.edited.srt", "人工修正後字幕", "主要字幕交付檔"],
            ["burn-settings.json", "結構化燒錄設定", "提供腳本讀取"],
            ["burn-settings.ffmpeg-style.txt", "FFmpeg force_style 字串", "人工檢查或手動 FFmpeg 使用"],
            ["export-manifest.json", "輸出摘要", "記錄 cue 數、修改數、警告數與輸出時間"],
            ["media_subtitled_sample_20s.mp4", "短樣片", "確認字幕大小與位置"],
            ["media_subtitled.mp4", "完整燒字幕影片", "最終影片輸出"],
        ],
    ),
    "錯誤防呆": (
        ["問題", "防呆方式"],
        [
            ["PowerShell 顯示亂碼", "使用 Python 以 UTF-8 讀檔驗證。"],
            ["Whisper cp950 錯誤", "設定 PYTHONIOENCODING=utf-8 與 PYTHONUTF8=1。"],
            ["PowerShell 不支援 Bash heredoc", "改用 PowerShell here-string 或獨立 .py 檔。"],
            ["字幕太大遮住 UI", "先輸出 20 秒樣片，再正式輸出。"],
            ["大型影片重複複製", "同磁碟可使用 Hard Link。"],
        ],
    ),
}


def write_markdown():
    lines = [
        "# Subtitle Review Loop 專案技術文件",
        "",
        "文件版本：1.0  ",
        "最後整理：2026-06-29  ",
        "Repository：`twyderek/claude.subtitle-review-loop`",
        "",
        "> 核心原則：AI 與腳本負責加速初稿與重複工作；人工校稿負責語意、專有名詞與畫面可讀性；所有輸出設定都必須可追溯、可重跑。",
        "",
    ]
    for heading, paragraphs in SECTIONS:
        lines.extend([f"## {heading}", ""])
        for paragraph in paragraphs:
            lines.extend([paragraph, ""])
        if heading.startswith("2."):
            append_md_table(lines, *TABLES["需求收斂"])
        if heading.startswith("5."):
            append_md_table(lines, *TABLES["模組清單"])
        if heading.startswith("6."):
            lines.extend([
                "```text",
                "workspace/media.mp4 -> workspace/media.rule-cleaned.srt -> browser review -> workspace/review-output/media.edited.srt -> burn-settings.json -> FFmpeg sample/full output",
                "```",
                "",
            ])
            append_md_table(lines, *TABLES["輸出格式"])
        if heading.startswith("7."):
            append_md_table(lines, *TABLES["錯誤防呆"])
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def append_md_table(lines, headers, rows):
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")


def make_diagram():
    width, height = 1800, 980
    img = Image.new("RGB", (width, height), "#0f172a")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("C:/Windows/Fonts/msjhbd.ttc", 42)
        heading_font = ImageFont.truetype("C:/Windows/Fonts/msjhbd.ttc", 24)
        body_font = ImageFont.truetype("C:/Windows/Fonts/msjh.ttc", 18)
        small_font = ImageFont.truetype("C:/Windows/Fonts/msjh.ttc", 15)
    except OSError:
        title_font = heading_font = body_font = small_font = ImageFont.load_default()

    def box(x, y, w, h, outline, title, body):
        draw.rounded_rectangle([x, y, x + w, y + h], radius=18, outline=outline, width=3)
        draw.text((x + 18, y + 15), title, fill="#e2e8f0", font=heading_font)
        yy = y + 54
        for line in body.split("\n"):
            draw.text((x + 18, yy), line, fill="#cbd5e1", font=small_font)
            yy += 24

    def arrow(x1, y1, x2, y2, color="#94a3b8"):
        draw.line([x1, y1, x2, y2], fill=color, width=3)
        pts = [(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)] if x2 >= x1 else [(x2, y2), (x2 + 12, y2 - 7), (x2 + 12, y2 + 7)]
        draw.polygon(pts, fill=color)

    draw.text((70, 48), "Subtitle Review Loop 技術工作流", fill="#e2e8f0", font=title_font)
    draw.text((72, 104), "從輸入、ASR/字幕清理、人工校稿，到校稿包與 FFmpeg 輸出", fill="#94a3b8", font=body_font)
    cards = [
        ("01 輸入資料", "影片 / SRT / rule.txt\n講義與專有名詞", "#3b82f6"),
        ("02 草稿字幕", "匯入既有 SRT\n或 ASR 初稿", "#8b5cf6"),
        ("03 規則清理", "只套用使用者\n提供的規則", "#22c55e"),
        ("04 階段一校稿", "看影片修字幕\n搜尋與跳 cue", "#f59e0b"),
        ("05 階段二預覽", "字型 / 大小 / 位置\n顏色 / 外框", "#06b6d4"),
        ("06 校稿包", "edited SRT\nsettings / manifest", "#ec4899"),
    ]
    xs = [70, 350, 630, 910, 1190, 1470]
    for i, (title, body, color) in enumerate(cards):
        box(xs[i], 170, 230, 128, color, title, body)
        if i:
            arrow(xs[i] - 42, 234, xs[i] - 8, 234)
    lower = [
        (140, 390, 360, 160, "#3b82f6", "前端工作台", "src/subtitle-editor.html\n兩階段 UI、SRT parse/build\n影片同步與字幕預覽"),
        (560, 390, 360, 160, "#22c55e", "本機服務", "src/subtitle-editor-server.mjs\n靜態檔案、MP4 Range\nPOST /api/save-review-package"),
        (980, 390, 360, 160, "#f59e0b", "輸出與燒錄", "scripts/burn-subtitles.mjs\n讀取 burn-settings.json\n產生樣片與正式 MP4"),
        (1400, 390, 300, 160, "#8b5cf6", "可攜工具包", "build-portable-toolkit.mjs\ninput / output\n雙擊啟動 CMD"),
    ]
    for item in lower:
        box(*item)
    for x1, x2 in [(500, 560), (920, 980), (1340, 1400)]:
        arrow(x1, 470, x2, 470)
    outputs = [
        (140, 650, 330, 130, "#06b6d4", "完整專案輸出", "workspace/review-output/\nmedia.edited.srt\nburn-settings.json"),
        (540, 650, 330, 130, "#ec4899", "可攜工具包輸出", "output/\nmedia.edited.srt\nexport-manifest.json"),
        (940, 650, 330, 130, "#a3e635", "樣片確認", "20 秒樣片先檢查\n字幕不遮擋 UI\n通過後才正式輸出"),
        (1340, 650, 330, 130, "#ef4444", "最終交付", "可編輯 SRT\n燒字幕 MP4\nQA Report"),
    ]
    for item in outputs:
        box(*item)
    arrow(1270, 715, 1340, 715)
    arrow(1120, 650, 1080, 595, "#f59e0b")
    draw.text((1010, 575), "未通過則回修樣式或字幕", fill="#fbbf24", font=small_font)
    draw.text((70, 890), "核心原則：自動化產生初稿，人工確認語意與畫面，所有設定留痕並可重跑。", fill="#fbbf24", font=body_font)
    img.save(DIAGRAM_PATH)


def set_east_asia(run):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_para(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_east_asia(run)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(9.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True)
        shade_cell(table.rows[0].cells[index], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    doc.add_paragraph()
    return table


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(31, 58, 95)
    return paragraph


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.text = "Subtitle Review Loop｜專案技術文件"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    footer = section.footer.paragraphs[0]
    footer.text = "Generated technical documentation"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    title = doc.add_paragraph()
    run = title.add_run("Subtitle Review Loop")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    set_east_asia(run)
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("專案技術文件")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(46, 116, 181)
    set_east_asia(run)
    add_para(doc, "文件版本：1.0｜整理日期：2026-06-29｜Repository：twyderek/claude.subtitle-review-loop")
    add_para(doc, "本文件整理專案需求、系統架構、資料流、模組設計、品質控管、部署方式與可攜工具包設計。")

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    shade_cell(callout.cell(0, 0), "F4F6F9")
    set_cell_text(callout.cell(0, 0), "核心原則：AI 與腳本負責加速初稿與重複工作；人工校稿負責語意、專有名詞與畫面可讀性；所有輸出設定都必須可追溯、可重跑。")
    doc.add_paragraph()
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.5))
    caption = doc.add_paragraph("圖 1：Subtitle Review Loop 技術工作流摘要")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()
    doc.add_heading("目錄", level=1)
    for heading, _ in SECTIONS:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(heading)
        set_east_asia(run)

    doc.add_page_break()
    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            add_para(doc, paragraph)
        if heading.startswith("2."):
            add_table(doc, *TABLES["需求收斂"])
        if heading.startswith("4."):
            for step in [
                "輸入收集：取得影片、草稿字幕、rule 檔、講義與專有名詞對照。",
                "字幕產生或匯入：使用既有 SRT，或以 Whisper/ASR 產生初稿。",
                "規則清理：依使用者提供的 rule 檔清理；沒有 rule 時只做中立檢查。",
                "字幕校稿：在瀏覽器中播放影片並修改字幕內容。",
                "燒錄預覽：設定字型、大小、位置、顏色、外框與粗體。",
                "校稿包輸出：輸出修正後 SRT、JSON 設定、FFmpeg style 與 manifest。",
                "樣片與正式輸出：先產生短樣片確認，再進行正式燒字幕。",
            ]:
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(step)
                set_east_asia(run)
        if heading.startswith("5."):
            add_table(doc, *TABLES["模組清單"])
        if heading.startswith("6."):
            add_code(doc, "workspace/media.mp4 -> workspace/media.rule-cleaned.srt -> browser review -> workspace/review-output/media.edited.srt -> burn-settings.json -> FFmpeg sample/full output")
            add_table(doc, *TABLES["輸出格式"])
        if heading.startswith("7."):
            add_table(doc, *TABLES["錯誤防呆"])
        if heading.startswith("8."):
            doc.add_heading("8.1 完整專案模式", level=2)
            add_code(doc, "npm install\nnpm run open")
            add_para(doc, "瀏覽器開啟：http://127.0.0.1:8787/src/subtitle-editor.html")
            doc.add_heading("8.2 Windows 雙擊啟動", level=2)
            add_code(doc, "start-subtitle-editor.cmd")
            doc.add_heading("8.3 可攜工具包", level=2)
            add_code(doc, "npm run build:toolkit")
            add_table(doc, ["完整專案", "可攜工具包"], [["workspace/media.mp4", "input/media.mp4"], ["workspace/media.rule-cleaned.srt", "input/media.srt"], ["workspace/review-output/", "output/"]])

    final = doc.add_table(rows=1, cols=1)
    final.style = "Table Grid"
    shade_cell(final.cell(0, 0), "E8EEF5")
    set_cell_text(final.cell(0, 0), "結論：Subtitle Review Loop 的價值，是把字幕製作流程拆成可檢查、可回溯、可重跑的模組。AI 產生初稿，規則負責一致化，人負責最後語意與畫面品質，FFmpeg 負責正式輸出。")
    doc.save(DOCX_PATH)


def main():
    write_markdown()
    make_diagram()
    build_docx()
    print(MD_PATH)
    print(DOCX_PATH)
    print(DIAGRAM_PATH)


if __name__ == "__main__":
    main()
